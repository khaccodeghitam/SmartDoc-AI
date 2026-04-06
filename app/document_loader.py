from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_core.documents import Document


def load_pdf(path: str | Path) -> List[Document]:
    loader = PDFPlumberLoader(str(path))
    return loader.load()


def load_docx(path: str | Path) -> List[Document]:
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    full_text = "\n".join(paragraphs).strip()

    if not full_text:
        return []

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": str(path),
                "file_type": "docx",
            },
        )
    ]


def load_documents(path: str | Path) -> List[Document]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf(file_path)
    if suffix == ".docx":
        return load_docx(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")
