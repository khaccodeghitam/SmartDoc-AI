"""PDF/DOCX document loading, saving, and raw text extraction."""
from __future__ import annotations

import io
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, List

from docx import Document as DocxDocument
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_core.documents import Document
from PIL import Image

from src.config import RAW_DIR
from src.utils import sanitize_name, source_name_from_path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def split_stuck_vietnamese_words(text: str) -> str:
    """
    Fixed heuristic to split only specific impossible combinations in Vietnamese PDFs.
    Avoids using character classes [] that cause over-splitting of common 'nh', 'ch'.
    """
    # Fix 'thủcông' -> 'thủ công', 'ngữcảnh' -> 'ngữ cảnh'
    text = text.replace('ủc', 'ủ c').replace('ữc', 'ữ c')
    # Fix 'địnhhướng' -> 'định hướng'
    text = text.replace('ịnhh', 'ịnh h')
    # Fix 'vựcưu' -> 'vực ưu'
    text = text.replace('ựcư', 'ực ư')
    # Fix 'mạnhmẽ' if stuck
    text = text.replace('ạnhm', 'ạnh m')
    return text


def _join_spans_with_spacing(spans: list[dict]) -> str:
    parts: list[str] = []
    prev: dict | None = None
    for span in spans:
        text = span.get("text", "")
        if not text:
            continue
        if prev is not None:
            prev_bbox = prev.get("bbox")
            curr_bbox = span.get("bbox")
            if prev_bbox and curr_bbox:
                gap = curr_bbox[0] - prev_bbox[2]
                size = min(prev.get("size", 10.0), span.get("size", 10.0))
                threshold = max(1.0, size * 0.2)
                if gap > threshold and parts and not parts[-1].endswith(" "):
                    parts.append(" ")
            else:
                if parts and parts[-1] and parts[-1][-1].isalnum() and text[0].isalnum():
                    parts.append(" ")
        parts.append(text)
        prev = span
    return "".join(parts)


def _score_extracted_text(text: str) -> float:
    if not text:
        return 1e9
    lines = [line.strip() for line in text.splitlines()]
    non_empty_lines = [line for line in lines if line]
    newline_ratio = text.count("\n") / max(1, len(text))

    short_lines = [line for line in non_empty_lines if len(line.split()) <= 3]
    short_line_ratio = len(short_lines) / max(1, len(non_empty_lines))

    tokens = re.findall(r"\S+", text)
    long_tokens = [token for token in tokens if len(token) >= 20]
    long_token_ratio = len(long_tokens) / max(1, len(tokens))

    camel_hits = len(re.findall(r"[a-zà-ỹ][A-ZÀ-Ỹ]", text))
    camel_ratio = camel_hits / max(1, len(text))

    return (2.0 * short_line_ratio) + (1.5 * newline_ratio) + (2.5 * long_token_ratio) + (3.0 * camel_ratio)


def _extract_pdfplumber_text(path: str | Path) -> tuple[str, list[Document]]:
    docs: list[Document] = []
    parts: list[str] = []
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text(
                    layout=True,
                    x_tolerance=2,
                    y_tolerance=2,
                )
                if text:
                    cleaned = text.strip().replace("</div>", "").replace("<div>", "")
                    if cleaned:
                        parts.append(cleaned)
                        docs.append(Document(page_content=cleaned, metadata={"source": str(path), "page": page_num}))
        return "\n\n".join(parts).strip(), docs
    except Exception:
        return "", []


# ---------------------------------------------------------------------------
# Document loading
# ---------------------------------------------------------------------------

def load_pdf_advanced(path: str | Path) -> List[Document]:
    """Advanced PDF loader with smart column sorting and OCR."""
    try:
        import fitz  # PyMuPDF
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        tessdata_dir = str(Path(__file__).resolve().parent.parent.parent / "tessdata")
        if os.path.exists(tessdata_dir):
            os.environ["TESSDATA_PREFIX"] = tessdata_dir

        full_text_parts = []
        pages_with_columns = 0
        abnormal_pages = 0
        prefer_column_first = False
        page_num = 0

        with fitz.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf, start=1):
                page_width = page.rect.width
                page_height = page.rect.height
                mid_x = page_width / 2

                row_pair_threshold = 0.4
                interleave_fallback_threshold = 0.55
                row_min_interleave_threshold = 0.15
                row_y_tolerance = max(6.0, page_height * 0.012)

                def get_block_group(b):
                    x0, y0, x1, y1 = b[:4]
                    # Titles/Full-width headers
                    if (x1 - x0) > page_width * 0.65:
                        return 0
                    # Left column
                    if x1 <= mid_x + 15:
                        return 1
                    # Right column
                    if x0 >= mid_x - 15:
                        return 2
                    return 0

                # Sử dụng cấu trúc Dictionary để lấy cả Text và Image theo đúng thứ tự
                page_dict = page.get_text("dict")
                page_elements = []

                for block in page_dict["blocks"]:
                    bbox = block["bbox"] # (x0, y0, x1, y1)
                    group = get_block_group(bbox)

                    if block["type"] == 0: # Khối văn bản
                        # Ghép các dòng trong block
                        block_text_lines: list[str] = []
                        for line in block["lines"]:
                            line_text = _join_spans_with_spacing(line.get("spans", [])).strip()
                            if line_text:
                                block_text_lines.append(line_text)

                        block_text = "\n".join(block_text_lines)

                        content = block_text.strip()
                        if content:
                            content = split_stuck_vietnamese_words(content)
                            page_elements.append({
                                "y0": bbox[1], "y1": bbox[3], "x": bbox[0], "x1": bbox[2],
                                "group": group, "content": content
                            })

                    elif block["type"] == 1: # Khối hình ảnh
                        try:
                            image_bytes = block["image"]
                            if image_bytes and len(image_bytes) > 5000:
                                img = Image.open(io.BytesIO(image_bytes))
                                ocr_res = pytesseract.image_to_string(img, lang="vie").strip()
                                if ocr_res and len(ocr_res) > 10:
                                    page_elements.append({
                                        "y0": bbox[1], "y1": bbox[3], "x": bbox[0], "x1": bbox[2],
                                        "group": group,
                                        "content": f"\n[Nội dung từ ảnh]:\n{ocr_res}"
                                    })
                        except: pass

                # Nếu trang trống trơn (có thể là trang scan hoàn toàn)
                if not page_elements:
                    try:
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        full_page_ocr = pytesseract.image_to_string(img, lang="vie").strip()
                        if len(full_page_ocr) > 20:
                            page_elements.append({
                                "y0": 0, "y1": page_height, "x": 0, "x1": page_width, "group": 0,
                                "content": f"\n[Nội dung quét toàn trang]:\n{full_page_ocr}"
                            })
                    except: pass

                # --- THUẬT TOÁN SẮP XẾP THEO VÙNG (REGIONAL COLUMN SORTING) ---
                # Sắp xếp tất cả theo Y để xử lý từ trên xuống dưới
                page_elements.sort(key=lambda e: e["y0"])

                left_blocks = [e for e in page_elements if e["group"] == 1]
                right_blocks = [e for e in page_elements if e["group"] == 2]
                top_fullwidth = any(
                    e["group"] == 0 and e["y0"] <= page_height * 0.22
                    for e in page_elements
                )
                gutter_ratio = 0.0
                if left_blocks and right_blocks:
                    left_max_x1 = max(e.get("x1", e["x"]) for e in left_blocks)
                    right_min_x0 = min(e["x"] for e in right_blocks)
                    gutter_ratio = (right_min_x0 - left_max_x1) / max(1.0, page_width)

                column_separated = gutter_ratio >= 0.04
                if (top_fullwidth or column_separated) and left_blocks and right_blocks:
                    if len(left_blocks) >= 3 and len(right_blocks) >= 3:
                        prefer_column_first = True

                def compute_row_pair_ratio(left, right) -> float:
                    if not left or not right:
                        return 0.0
                    pairs = 0
                    used_right = set()
                    for l in left:
                        l_height = max(1.0, l["y1"] - l["y0"])
                        for idx, r in enumerate(right):
                            if idx in used_right:
                                continue
                            r_height = max(1.0, r["y1"] - r["y0"])
                            overlap = min(l["y1"], r["y1"]) - max(l["y0"], r["y0"])
                            if overlap >= 0.4 * min(l_height, r_height) or abs(l["y0"] - r["y0"]) <= row_y_tolerance:
                                pairs += 1
                                used_right.add(idx)
                                break
                    return pairs / max(1, min(len(left), len(right)))

                def compute_interleave_ratio(elements) -> float:
                    groups = [e["group"] for e in elements if e["group"] in (1, 2)]
                    if len(groups) < 2:
                        return 0.0
                    switches = sum(1 for i in range(1, len(groups)) if groups[i] != groups[i - 1])
                    return switches / (len(groups) - 1)

                row_pair_ratio = compute_row_pair_ratio(left_blocks, right_blocks)
                layout_mode = "row" if row_pair_ratio >= row_pair_threshold else "column"
                if (top_fullwidth or column_separated) and left_blocks and right_blocks:
                    if len(left_blocks) >= 3 and len(right_blocks) >= 3:
                        layout_mode = "column"

                if layout_mode == "row":
                    final_page_elements = sorted(page_elements, key=lambda e: (e["y0"], e["x"]))
                else:
                    final_page_elements = []
                    temp_column_elements = []

                    for el in page_elements:
                        if el["group"] == 0:
                            # Nếu gặp Group 0 (Toàn khổ), ta phải "xả" hết các cột đang chờ
                            if temp_column_elements:
                                # Sắp xếp các cột đang chờ: Nhóm 1 (Trái) trước, rồi đến Nhóm 2 (Phải)
                                temp_column_elements.sort(key=lambda x: (x["group"], x["y0"]))
                                final_page_elements.extend(temp_column_elements)
                                temp_column_elements = []
                            final_page_elements.append(el)
                        else:
                            # Gom các khối cột lại để xử lý sau
                            temp_column_elements.append(el)

                    # Xả nốt các khối cột còn lại ở cuối trang
                    if temp_column_elements:
                        temp_column_elements.sort(key=lambda x: (x["group"], x["y0"]))
                        final_page_elements.extend(temp_column_elements)

                if left_blocks and right_blocks:
                    pages_with_columns += 1
                    interleave_ratio = compute_interleave_ratio(final_page_elements)
                    if layout_mode == "column" and interleave_ratio >= interleave_fallback_threshold:
                        abnormal_pages += 1
                    elif layout_mode == "row" and row_pair_ratio >= row_pair_threshold and interleave_ratio <= row_min_interleave_threshold:
                        abnormal_pages += 1

                for el in final_page_elements:
                    full_text_parts.append(el["content"])

        raw_text_pymupdf = "\n\n".join(full_text_parts).strip().replace("</div>", "").replace("<div>", "")
        pymupdf_score = _score_extracted_text(raw_text_pymupdf)

        plumber_text, _ = _extract_pdfplumber_text(path)
        plumber_score = _score_extracted_text(plumber_text)

        force_plumber = False
        if pages_with_columns:
            abnormal_ratio = abnormal_pages / pages_with_columns
            if abnormal_ratio >= 0.3:
                force_plumber = True

        use_plumber = False
        if prefer_column_first and raw_text_pymupdf and not force_plumber:
            use_plumber = False
        else:
            if plumber_text and (force_plumber or plumber_score < pymupdf_score):
                use_plumber = True
            if not raw_text_pymupdf and plumber_text:
                use_plumber = True

        selected_text = plumber_text if use_plumber else raw_text_pymupdf
        extraction_method = "pdfplumber_layout_auto" if use_plumber else "fitz_regional_auto"

        # Dùng \n\n để phân tách các khối rõ ràng, giúp Splitter nhận diện đoạn văn tốt hơn
        final_text = selected_text
        # Loại bỏ các dấu xuống dòng thừa thãi trong nội dung để tránh split vụn
        final_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', final_text)
        return [Document(page_content=final_text, metadata={"source": str(path), "file_type": "pdf", "page": page_num, "extraction_method": extraction_method})]
    except Exception as e:
        print(f"⚠️ PDF Advanced failed: {e}")
        return load_pdf(path)


def load_docx(path: str | Path) -> List[Document]:
    """Context-aware DOCX loader."""
    try:
        doc = DocxDocument(str(path))
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        tessdata_dir = str(Path(__file__).resolve().parent.parent.parent / "tessdata")
        if os.path.exists(tessdata_dir):
            os.environ["TESSDATA_PREFIX"] = tessdata_dir
        full_content_parts = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_content_parts.append(text)
            
            for run in paragraph.runs:
                if 'drawing' in run.element.xml:
                    try:
                        blip_ids = re.findall(r'r:embed="([^"]+)"', run.element.xml)
                        for rId in blip_ids:
                            image_part = doc.part.related_parts[rId]
                            img = Image.open(io.BytesIO(image_part.blob))
                            ocr_res = pytesseract.image_to_string(img, lang="vie").strip()
                            if ocr_res and len(ocr_res) > 15:
                                full_content_parts.append(f"\n[Nội dung hình ảnh tại đây]:\n{ocr_res}\n")
                    except: pass

        all_text = "\n".join(full_content_parts).strip()
        # No splitting logic for Word as it is usually clean
        return [Document(page_content=all_text, metadata={"source": str(path), "file_type": "docx"})]
    except Exception as e:
        print(f"⚠️ load_docx failed: {e}")
        doc = DocxDocument(str(path))
        return [Document(page_content="\n".join([p.text for p in doc.paragraphs]), metadata={"source": str(path)})]


def load_pdf(path: str | Path) -> List[Document]:
    """Fallback loader."""
    _, docs = _extract_pdfplumber_text(path)
    return docs

def load_documents(path: str | Path, use_advanced_pdf: bool = False) -> List[Document]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf_advanced(file_path) if use_advanced_pdf else load_pdf(file_path)
    if suffix == ".docx":
        return load_docx(file_path)
    raise ValueError(f"Unsupported: {suffix}")

def save_uploaded_file(uploaded_file: Any, target_dir: Path = RAW_DIR) -> Path:
    target_dir.mkdir(parents=True, exist_ok=True)
    file_path = target_dir / uploaded_file.name
    file_path.write_bytes(uploaded_file.getbuffer())
    return file_path

def enrich_chunks_metadata(chunks: list[Document], file_path: Path) -> list[Document]:
    source = str(file_path)
    for doc in chunks:
        m = dict(doc.metadata or {})
        m.setdefault("source", source)
        m["file_name"] = file_path.name
        m["file_type"] = file_path.suffix.lower().lstrip(".")
        m["upload_time"] = datetime.now().isoformat(timespec="seconds")
        m["upload_date"] = m["upload_time"][:10]
        doc.metadata = m
    return chunks
